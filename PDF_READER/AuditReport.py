import os
from docx import Document
import collections
from pdf_extractor import PDFExtractor
from keywords_checker import KeywordChecker
from BarPlotVisualizer import BarPlotVisualizer
from datetime import datetime,date


#______________Vars
TODAY= date.today()
TODAY_FORMATED = TODAY.strftime("%B %d, %Y")


class AuditReport:
    """
    Encapsulates the processing of each PDF into its own method (_process_pdf).
    Separates the concerns of printing to the console and adding to the document into their own methods (_print_to_console and _add_to_document, respectively).
    Simplifies the main flow in the generate_report method.
    """

    def __init__(self, dir_path, keywords_filepath, doc_filepath, route_name):
        self.route_name = route_name
        self.dir_path = dir_path
        self.doc_filepath = doc_filepath
        self.keywords_filepath = keywords_filepath
        self.doc = Document()
        self.keyword_checker = KeywordChecker(self.keywords_filepath)
        self.client_keyword_counts = collections.defaultdict(int)

    def generate_report(self):
        pdf_files = [file for file in os.listdir(self.dir_path) if file.endswith('.pdf')]

        for pdf_file in pdf_files:
            pdf_filepath = os.path.join(self.dir_path, pdf_file)
            pdf_extractor = PDFExtractor(pdf_filepath)

            self._process_pdf(pdf_extractor, pdf_file)


        DOC_FILENAME = self.doc_filepath
        self.doc.save(DOC_FILENAME)

        visualizer = BarPlotVisualizer(self.client_keyword_counts, 'Client Names', 'Number of Safeguarding Keywords',
                                       f'Number of Safeguarding Keywords Found per Client as of {TODAY_FORMATED} for {self.route_name}')
        visualizer.plot()

        print(self.client_keyword_counts)

    def _process_pdf(self, pdf_extractor, pdf_file):
        client_name = pdf_extractor.extract_client_name()
        carer_name = pdf_extractor.extract_carer_name()
        check_in_time = pdf_extractor.extract_check_in_time()
        notes = pdf_extractor.extract_notes()

        found_keywords = self.keyword_checker.keywords_found(notes)
        self.client_keyword_counts[client_name] += len(found_keywords)

        notes_with_emoji = self.keyword_checker.replace_keywords_with_emoji(notes)

        if found_keywords:
            self._print_to_console(pdf_file, client_name, carer_name, check_in_time, notes_with_emoji, found_keywords)
            self._add_to_document(pdf_file, client_name, carer_name, check_in_time, notes_with_emoji, found_keywords)

    def _print_to_console(self, pdf_file, client_name, carer_name, check_in_time, notes_with_emoji, found_keywords):
        print(f"""
            PDF: {pdf_file}
            Birdie Care Log for {client_name},
            Visit by {carer_name},
            Check-in: {check_in_time},
            Notes: "{notes_with_emoji}"
        """)
        print("Keywords found:", ", ".join(found_keywords))
        print(f"⛔{client_name} experienced,", "".join(found_keywords))
        print(f"according to {carer_name}")
        print("----------------------------------------------------------------------")

    def _add_to_document(self, pdf_file, client_name, carer_name, check_in_time, notes_with_emoji, found_keywords):
        self.doc.add_paragraph(f"PDF: {pdf_file}")
        self.doc.add_paragraph(f"Birdie Care Log for {client_name},")
        self.doc.add_paragraph(f"Visit by {carer_name},")
        self.doc.add_paragraph(f"Check-in: {check_in_time},")
        self.doc.add_paragraph(f"Notes: \"{notes_with_emoji}\"")
        self.doc.add_paragraph(f"Keywords found: {', '.join(found_keywords)}")
        self.doc.add_paragraph(f"⛔{client_name} experienced, {', '.join(found_keywords)}")
        self.doc.add_paragraph(f"according to {carer_name}")
        self.doc.add_paragraph("----------------------------------------------------------------------")


